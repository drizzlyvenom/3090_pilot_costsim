from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph


PACKAGE_ROOT = Path(__file__).resolve().parents[2]
RUN_DIR = PACKAGE_ROOT / "02_simulation_run"
PAPERS_DIR = PACKAGE_ROOT / "03_results_summary" / "short_papers"
FIGURES_DIR = PACKAGE_ROOT / "03_results_summary" / "figures"
RESULTS_DIR = RUN_DIR / "results"
RENDER_DIR = RUN_DIR / "render_check"


def read_csv(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_manifest() -> dict:
    return json.loads((RUN_DIR / "manifest.json").read_text(encoding="utf-8"))


def pct(old: float, new: float) -> float:
    return (old - new) / old * 100.0


def get_metrics() -> dict:
    primary = read_csv(RESULTS_DIR / "primary_case_summary.csv")
    manifest = read_manifest()
    by = {row["baseline"]: row for row in primary}
    b1 = by["B1"]
    b2 = by["B2"]
    b4 = by["B4-lite"]
    b5 = by["B5-lite"]
    b7 = by["B7-lite"]
    interpretation = manifest["interpretation"]
    return {
        "primary": primary,
        "b1_peak": float(b1["estimated_peak_memory_gb"]),
        "b2_peak": float(b2["estimated_peak_memory_gb"]),
        "b7_peak": float(b7["estimated_peak_memory_gb"]),
        "b1_tokens": int(b1["visual_tokens"]),
        "b2_tokens": int(b2["visual_tokens"]),
        "b4_adapter_gb": float(b4["resident_adapter_memory_mb"]) / 1024.0,
        "b5_adapter_gb": float(b5["resident_adapter_memory_mb"]) / 1024.0,
        "b7_adapter_gb": float(b7["resident_adapter_memory_mb"]) / 1024.0,
        "token_reduction": interpretation["token_reduction_b2"],
        "peak_reduction": interpretation["peak_reduction_b7"],
        "adapter_reduction_b5": interpretation["adapter_reduction_b5"],
        "adapter_reduction_b7": interpretation["adapter_reduction_b7"],
        "b4_fail_runs": interpretation["b4_fail_runs"],
        "b7_fail_runs": interpretation["b7_fail_runs"],
        "b7_adjusted_runs": interpretation["b7_adjusted_runs"],
        "b7_total_runs": interpretation["b7_total_runs"],
    }


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(9.4)

    for style_name, size in [("Heading 1", 13), ("Heading 2", 11)]:
        styles[style_name].font.name = "Malgun Gothic"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        styles[style_name].font.size = Pt(size)


def add_para(doc: Document, text: str, style: str | None = None, align: int | None = None, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.04
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, bold=True)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.02
        run = p.add_run(item)
        set_run_font(run, size=9.2)


def add_result_table(doc: Document, rows: list[dict], lang: str) -> None:
    headers = {
        "en": ["Baseline", "Tokens", "Peak GB", "Adapter GB", "Reserve", "Latency"],
        "ko": ["비교군", "토큰", "Peak GB", "Adapter GB", "Reserve", "Latency"],
    }[lang]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "E8EEF7")
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=8.4, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows:
        cells = table.add_row().cells
        values = [
            row["baseline"],
            row["visual_tokens"],
            f"{float(row['estimated_peak_memory_gb']):.2f}",
            f"{float(row['resident_adapter_memory_mb']) / 1024.0:.2f}",
            "pass" if row["reserve_pass"] == "True" else "fail",
            f"{float(row['latency_proxy_ms']):.1f}",
        ]
        for i, value in enumerate(values):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=8.2)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=8.2, italic=True)


def add_references(doc: Document, lang: str) -> None:
    refs_en = [
        "Hu, E. J. et al. (2021). LoRA: Low-Rank Adaptation of Large Language Models. arXiv:2106.09685.",
        "Tian, C. et al. (2024). HydraLoRA: An Asymmetric LoRA Architecture for Efficient Fine-Tuning. arXiv:2404.19245.",
        "Maes, L. et al. (2026). LeWorldModel: Stable End-to-End Joint-Embedding Predictive Architecture from Pixels. arXiv:2603.19312.",
        "Davidson, T. R. et al. (2026). Reasoning-Driven Synthetic Data Generation and Evaluation. arXiv:2603.29791.",
    ]
    refs_ko = [
        "Hu, E. J. 외 (2021). LoRA: Low-Rank Adaptation of Large Language Models. arXiv:2106.09685.",
        "Tian, C. 외 (2024). HydraLoRA: An Asymmetric LoRA Architecture for Efficient Fine-Tuning. arXiv:2404.19245.",
        "Maes, L. 외 (2026). LeWorldModel: Stable End-to-End Joint-Embedding Predictive Architecture from Pixels. arXiv:2603.19312.",
        "Davidson, T. R. 외 (2026). Reasoning-Driven Synthetic Data Generation and Evaluation. arXiv:2603.29791.",
    ]
    for ref in (refs_en if lang == "en" else refs_ko):
        add_para(doc, ref)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def make_en(metrics: dict) -> tuple[str, list[str]]:
    title = "A Cost-Simulation Feasibility Study of Budget-Conditioned Foveated Adapter Residency under 24GB GPU Memory"
    md: list[str] = [f"# {title}", ""]
    doc = Document()
    configure_document(doc)

    add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(doc, "Preliminary short-paper draft for research discussion", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_heading(doc, "Abstract")
    add_para(
        doc,
        "Physical-AI systems often require perception, planning, verification, and world-model components to share a limited GPU memory budget. This note evaluates whether a budget-conditioned foveated adapter residency policy is worth pursuing before implementing full VLM training or HydraLoRA-style fine-tuning. We build a cost simulator for an RTX 3090 24GB environment and compare six baselines: low-resolution global input, full-resolution input, foveated ROI input, independent adapter residency, shared adapter residency, and a budget-conditioned policy. Across 4,032 simulated configurations, the foveated input reduced visual tokens by 75.0% in the primary case, while the final budgeted variant reduced estimated peak memory by 21.9% relative to full-resolution input. The shared-bank and budgeted variants reduced resident adapter memory by 67.36% and 89.24%, respectively, compared with an independent bank. These results do not establish task accuracy, but they provide a compact feasibility signal for requesting resources for real VLM profiling and LoRA loading experiments.",
    )
    add_heading(doc, "Keywords", 2)
    add_para(doc, "foveated inference; adapter residency; LoRA; GPU memory; cost simulation; physical AI")
    page_break(doc)

    add_heading(doc, "1. Introduction")
    add_para(
        doc,
        "A deployment-oriented physical-AI loop is not only a question of running the largest possible vision model. In a practical stack, a perception model must leave memory for planning, verification, orchestration, and possibly a world model. This motivates a policy view: the system should decide where to look, which lightweight adapter to activate, and when to keep or evict that adapter under a memory budget.",
    )
    add_para(
        doc,
        "The proposed direction combines three ideas. First, foveated visual reasoning reduces input cost by starting with a low-resolution global observation and requesting high-resolution evidence only for selected regions. Second, LoRA-style adapters represent local skills that may be loaded selectively. Third, HydraLoRA-style sharing suggests that adapter banks need not scale linearly with the number of skills. The central question of this note is therefore modest: before expensive model training, does the structure produce a measurable memory-cost advantage in simulation?",
    )
    add_heading(doc, "Contributions", 2)
    add_bullets(
        doc,
        [
            "We define a small cost model that separates visual-token cost, adapter resident memory, temporary load buffers, and orchestration reserve.",
            "We compare full-resolution inference against foveated ROI inference and independent adapter banks against shared adapter banks.",
            "We report an initial stress-grid result showing that budget-conditioned residency can eliminate reserve failures in this simulator.",
        ],
    )
    page_break(doc)

    add_heading(doc, "2. Method")
    add_para(
        doc,
        "The simulator uses a proxy 7B VLM memory profile on an RTX 3090 24GB GPU. The primary case assumes a 12GB base model, a 4GB orchestration reserve target, a 1344 x 1344 full-resolution input, a 336 x 336 global view, three 336 x 336 ROI crops, sixteen candidate adapters, and top-k=2 active adapters. Visual tokens are estimated from a 14-pixel patch size. Memory is modeled as base model memory plus visual activation/KV proxy cost, resident adapter memory, temporary load buffer, and reserve constraints.",
    )
    add_para(
        doc,
        "Six baselines are evaluated. B0 uses only the low-resolution global view. B1 uses full-resolution input and serves as the expensive reference. B2 adds ROI crops without adapters. B4-lite assumes an independent adapter bank in which all candidate adapter weights are resident. B5-lite assumes a shared common component plus per-skill branches. B7-lite adds a rule-based budget controller that reduces top-k, reduces ROI count, or holds adapter loading when the reserve target would be violated.",
    )
    add_para(
        doc,
        "The evaluation grid spans three model profiles, three reserve targets, multiple input resolutions, ROI counts, adapter counts, and active top-k values. The output metrics are visual-token count, estimated peak memory, resident adapter memory, reserve pass/fail, load/evict/hold counts, and a latency proxy. Because this is a cost simulation, no accuracy claim is made.",
    )
    page_break(doc)

    add_heading(doc, "3. Results")
    add_para(
        doc,
        f"In the primary case, B1 full-resolution input used {metrics['b1_tokens']} visual tokens and {metrics['b1_peak']:.2f}GB estimated peak memory. B2 foveated ROI input used {metrics['b2_tokens']} tokens, a {metrics['token_reduction']:.1f}% token reduction. B7-lite reached {metrics['b7_peak']:.2f}GB estimated peak memory, a {metrics['peak_reduction']:.1f}% reduction relative to B1.",
    )
    add_result_table(doc, metrics["primary"], "en")
    add_caption(doc, "Table 1. Primary-case simulation results.")
    doc.add_picture(str(FIGURES_DIR / "primary_peak_memory_bars.png"), width=Inches(4.9))
    add_caption(doc, "Figure 1. Estimated peak memory by baseline in the primary case.")
    add_para(
        doc,
        f"Adapter residency shows the clearest structural separation. B4-lite used {metrics['b4_adapter_gb']:.2f}GB resident adapter memory, B5-lite used {metrics['b5_adapter_gb']:.2f}GB, and B7-lite used {metrics['b7_adapter_gb']:.2f}GB. In the full grid, B4-lite produced {metrics['b4_fail_runs']} reserve failures, whereas B7-lite produced {metrics['b7_fail_runs']} reserve failures and adjusted {metrics['b7_adjusted_runs']} of {metrics['b7_total_runs']} runs through top-k, ROI, or adapter-hold decisions.",
    )
    page_break(doc)

    add_heading(doc, "4. Discussion")
    add_para(
        doc,
        "The result supports the feasibility of separating two resource-control levers: foveated input selection reduces visual-token and activation-related cost, while shared or budgeted adapter residency reduces memory tied to skill specialization. The simulation also shows why a policy formulation is useful. A naive independent bank can pass in easy settings but fail under larger reserve targets or heavier model profiles. A budget controller can trade adapter breadth or ROI count for reserve safety.",
    )
    add_heading(doc, "Limitations", 2)
    add_para(
        doc,
        "The simulator is not a profiler. It does not measure kernel-level memory allocation, real VLM attention behavior, adapter loading overhead, or task accuracy. The adapter costs are first-pass assumptions and must be replaced with parameter-count-derived or profiler-derived measurements once a target VLM and LoRA rank are selected. Therefore, the present result should be treated as a feasibility note, not as a performance claim.",
    )
    add_heading(doc, "Conclusion and Next Step", 2)
    add_para(
        doc,
        "The 24GB GPU cost simulation indicates that budget-conditioned foveated adapter residency is worth validating with real model runs. The immediate next step is to profile a small VLM under full-resolution and foveated inputs, then replace the simulated adapter table with measured LoRA loading and residency costs. If supported by additional compute resources, the same protocol can be extended to real task accuracy and co-resident physical-AI modules.",
    )
    add_heading(doc, "References", 2)
    add_references(doc, "en")

    md.extend([
        "Abstract: " + doc.paragraphs[3].text,
        "",
        "Sections: Introduction, Method, Results, Discussion, References.",
        "",
        f"Key numbers: token reduction {metrics['token_reduction']}%, peak reduction {metrics['peak_reduction']}%, B4 failures {metrics['b4_fail_runs']}, B7 failures {metrics['b7_fail_runs']}.",
    ])
    return title, [doc, "\n".join(md)]


def make_ko(metrics: dict) -> tuple[str, list[str]]:
    title = "24GB GPU 메모리 제약에서의 Budget-Conditioned Foveated Adapter Residency 비용 시뮬레이션 연구"
    md: list[str] = [f"# {title}", ""]
    doc = Document()
    configure_document(doc)

    add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(doc, "후속 연구 논의를 위한 1차 소논문 초안", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_heading(doc, "초록")
    add_para(
        doc,
        "피지컬 AI 시스템에서는 perception, planning, verification, world model 계층이 제한된 GPU 메모리를 함께 사용해야 한다. 본 초안은 실제 VLM 학습이나 HydraLoRA 계열 fine-tuning을 수행하기 전에, budget-conditioned foveated adapter residency 정책이 후속 연구 가치가 있는지 비용 시뮬레이션으로 확인한다. RTX 3090 24GB 환경을 기준으로 cost simulator를 만들고, low-resolution global input, full-resolution input, foveated ROI input, independent adapter residency, shared adapter residency, budget-conditioned policy의 여섯 비교군을 평가했다. 총 4,032개 조건에서 대표 조건의 foveated input은 visual token을 75.0% 줄였고, 최종 B7-lite 구조는 full-resolution 기준선 대비 estimated peak memory를 21.9% 줄였다. 또한 shared-bank와 budgeted variant는 independent bank 대비 resident adapter memory를 각각 67.36%, 89.24% 줄였다. 이 결과는 task accuracy를 증명하지는 않지만, 실제 VLM profiling과 LoRA loading 실험을 위한 초기 가능성 근거를 제공한다.",
    )
    add_heading(doc, "핵심어", 2)
    add_para(doc, "foveated inference; adapter residency; LoRA; GPU memory; cost simulation; physical AI")
    page_break(doc)

    add_heading(doc, "1. 서론")
    add_para(
        doc,
        "배치 관점의 피지컬 AI 루프는 가장 큰 비전 모델 하나를 띄우는 문제가 아니다. 실제 시스템에서는 perception 모델이 planner, verifier, orchestration layer, world model이 사용할 메모리 여유를 남겨야 한다. 따라서 이 연구는 어디를 볼지, 어떤 lightweight adapter를 활성화할지, 그리고 언제 유지하거나 내릴지를 메모리 예산 아래에서 결정하는 정책 문제로 접근한다.",
    )
    add_para(
        doc,
        "제안 방향은 세 가지 생각을 결합한다. 첫째, foveated visual reasoning은 저해상도 전역 관측에서 시작하고 필요한 영역에 대해서만 고해상도 증거를 요청함으로써 입력 비용을 줄인다. 둘째, LoRA 계열 adapter는 국소 skill을 선택적으로 올릴 수 있는 단위로 볼 수 있다. 셋째, HydraLoRA식 공유 구조는 adapter bank의 메모리가 skill 수에 완전히 선형으로 증가하지 않을 가능성을 준다. 본 초안의 질문은 작다. 비싼 모델 학습 전에, 이 구조가 시뮬레이션 상에서 측정 가능한 메모리 비용 이득을 만드는가를 본다.",
    )
    add_heading(doc, "기여", 2)
    add_bullets(
        doc,
        [
            "visual-token cost, adapter resident memory, temporary load buffer, orchestration reserve를 분리한 작은 cost model을 정의했다.",
            "full-resolution inference와 foveated ROI inference, independent adapter bank와 shared adapter bank를 비교했다.",
            "budget-conditioned residency가 simulator 안에서 reserve failure를 줄일 수 있음을 stress grid로 확인했다.",
        ],
    )
    page_break(doc)

    add_heading(doc, "2. 방법")
    add_para(
        doc,
        "시뮬레이터는 RTX 3090 24GB 위의 proxy 7B VLM memory profile을 사용한다. 대표 조건은 base model 12GB, orchestration reserve target 4GB, full-resolution 1344 x 1344, global view 336 x 336, ROI crop 336 x 336 세 개, 후보 adapter 16개, active top-k=2로 둔다. visual token은 14픽셀 patch size로 추정한다. 전체 메모리는 base model memory, visual activation/KV proxy cost, resident adapter memory, temporary load buffer, reserve constraint로 계산한다.",
    )
    add_para(
        doc,
        "비교군은 여섯 개다. B0는 저해상도 global view만 쓴다. B1은 full-resolution 입력을 쓰는 비용 기준선이다. B2는 adapter 없이 ROI crop을 추가한다. B4-lite는 independent adapter bank가 모두 상주한다고 가정한다. B5-lite는 shared common component와 skill별 branch를 가정한다. B7-lite는 reserve target을 넘을 위험이 있으면 top-k를 줄이고, ROI 수를 줄이거나, adapter load를 hold하는 규칙 기반 budget controller를 추가한다.",
    )
    add_para(
        doc,
        "평가 grid는 세 가지 model profile, 세 가지 reserve target, 여러 해상도, ROI 수, adapter 수, active top-k 값을 포함한다. 측정값은 visual-token count, estimated peak memory, resident adapter memory, reserve pass/fail, load/evict/hold count, latency proxy다. 본 실험은 cost simulation이므로 정확도 주장은 하지 않는다.",
    )
    page_break(doc)

    add_heading(doc, "3. 결과")
    add_para(
        doc,
        f"대표 조건에서 B1 full-resolution 입력은 {metrics['b1_tokens']} visual token과 {metrics['b1_peak']:.2f}GB estimated peak memory를 사용했다. B2 foveated ROI 입력은 {metrics['b2_tokens']} token을 사용해 token 수를 {metrics['token_reduction']:.1f}% 줄였다. B7-lite는 {metrics['b7_peak']:.2f}GB estimated peak memory를 기록했으며, 이는 B1 대비 {metrics['peak_reduction']:.1f}% 감소다.",
    )
    add_result_table(doc, metrics["primary"], "ko")
    add_caption(doc, "표 1. 대표 조건의 시뮬레이션 결과.")
    doc.add_picture(str(FIGURES_DIR / "primary_peak_memory_bars.png"), width=Inches(4.9))
    add_caption(doc, "그림 1. 대표 조건에서 비교군별 estimated peak memory.")
    add_para(
        doc,
        f"adapter residency에서는 구조적 차이가 더 분명하다. B4-lite는 resident adapter memory {metrics['b4_adapter_gb']:.2f}GB를 사용했고, B5-lite는 {metrics['b5_adapter_gb']:.2f}GB, B7-lite는 {metrics['b7_adapter_gb']:.2f}GB를 사용했다. 전체 grid에서 B4-lite는 reserve failure {metrics['b4_fail_runs']}건을 보였지만, B7-lite는 reserve failure {metrics['b7_fail_runs']}건을 유지했고 {metrics['b7_total_runs']}건 중 {metrics['b7_adjusted_runs']}건에서 top-k, ROI, adapter hold 조정을 수행했다.",
    )
    page_break(doc)

    add_heading(doc, "4. 논의")
    add_para(
        doc,
        "결과는 두 개의 자원 제어 축을 분리해 볼 수 있음을 보여준다. foveated input selection은 visual token과 activation 계열 비용을 줄이고, shared 또는 budgeted adapter residency는 skill specialization에 묶인 resident memory를 줄인다. 또한 policy formulation이 필요한 이유도 드러난다. naive independent bank는 쉬운 조건에서는 통과하지만, reserve target이 커지거나 model profile이 무거워질수록 실패한다. 반면 budget controller는 adapter 폭이나 ROI 수를 조정해 reserve safety를 확보한다.",
    )
    add_heading(doc, "한계", 2)
    add_para(
        doc,
        "이 시뮬레이터는 profiler가 아니다. 실제 kernel-level memory allocation, VLM attention behavior, adapter loading overhead, task accuracy를 측정하지 않는다. adapter cost도 1차 가정값이므로 target VLM과 LoRA rank가 정해지면 parameter count 또는 profiler 기반 측정값으로 교체해야 한다. 따라서 본 결과는 성능 주장이라기보다 feasibility note로 해석해야 한다.",
    )
    add_heading(doc, "결론과 다음 단계", 2)
    add_para(
        doc,
        "24GB GPU cost simulation은 budget-conditioned foveated adapter residency가 실제 모델 실험으로 검증할 가치가 있음을 보여준다. 다음 단계는 작은 VLM을 대상으로 full-resolution 입력과 foveated 입력의 실제 peak memory를 측정하고, simulated adapter table을 실제 LoRA loading/residency cost로 교체하는 것이다. 추가 compute 지원을 받으면 동일 프로토콜을 task accuracy와 co-resident physical-AI module 검증으로 확장할 수 있다.",
    )
    add_heading(doc, "참고문헌", 2)
    add_references(doc, "ko")

    md.extend([
        "초록: " + doc.paragraphs[3].text,
        "",
        "구성: 서론, 방법, 결과, 논의, 참고문헌.",
        "",
        f"핵심 수치: token reduction {metrics['token_reduction']}%, peak reduction {metrics['peak_reduction']}%, B4 failures {metrics['b4_fail_runs']}, B7 failures {metrics['b7_fail_runs']}.",
    ])
    return title, [doc, "\n".join(md)]


def verify_docx(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return {
        "path": str(path),
        "readable": True,
        "paragraphs": len(paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "first_paragraph": paragraphs[0] if paragraphs else "",
    }


def iter_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, doc)


def docx_to_markdown(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []
    for block in iter_blocks(doc):
        if isinstance(block, DocxParagraph):
            text = block.text.strip()
            if not text:
                continue
            style = block.style.name if block.style else ""
            if style == "Heading 1":
                lines.extend([f"## {text}", ""])
            elif style == "Heading 2":
                lines.extend([f"### {text}", ""])
            elif style.startswith("List Bullet"):
                lines.append(f"- {text}")
                lines.append("")
            elif not lines:
                lines.extend([f"# {text}", ""])
            else:
                lines.extend([text, ""])
        elif isinstance(block, DocxTable):
            rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in block.rows]
            if rows:
                lines.append("| " + " | ".join(rows[0]) + " |")
                lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
                for row in rows[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
    return "\n".join(lines).strip() + "\n"


def main() -> int:
    PAPERS_DIR.mkdir(parents=True, exist_ok=True)
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    metrics = get_metrics()

    _, en_payload = make_en(metrics)
    en_doc, en_md = en_payload
    _, ko_payload = make_ko(metrics)
    ko_doc, ko_md = ko_payload

    en_docx = PAPERS_DIR / "short_paper_costsim_en.docx"
    ko_docx = PAPERS_DIR / "short_paper_costsim_ko.docx"
    en_src = PAPERS_DIR / "short_paper_costsim_en.md"
    ko_src = PAPERS_DIR / "short_paper_costsim_ko.md"

    en_doc.save(en_docx)
    ko_doc.save(ko_docx)
    en_src.write_text(docx_to_markdown(en_docx), encoding="utf-8")
    ko_src.write_text(docx_to_markdown(ko_docx), encoding="utf-8")

    check = {
        "english": verify_docx(en_docx),
        "korean": verify_docx(ko_docx),
        "note": "DOCX files include explicit page breaks for a five-section short-paper draft. Exact pagination should be confirmed in Word before PDF export.",
    }
    (RENDER_DIR / "short_paper_docx_check.json").write_text(json.dumps(check, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(check, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
