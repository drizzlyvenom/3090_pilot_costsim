from __future__ import annotations

import csv
import datetime as dt
import json
import math
import os
import statistics
import subprocess
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


PACKAGE_ROOT = Path(__file__).resolve().parents[2]
RUN_DIR = PACKAGE_ROOT / "02_simulation_run"
CONFIG_DIR = RUN_DIR / "config"
RESULTS_DIR = RUN_DIR / "results"
FIGURES_DIR = PACKAGE_ROOT / "03_results_summary" / "figures"
REPORTS_DIR = PACKAGE_ROOT / "03_results_summary" / "technical_report"
RENDER_DIR = RUN_DIR / "render_check"

MALGUN = Path(r"C:\Windows\Fonts\malgun.ttf")
MALGUN_BOLD = Path(r"C:\Windows\Fonts\malgunbd.ttf")


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def write_json(path: Path, data: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def tokens_for_resolution(resolution: list[int], patch_size: int) -> int:
    width, height = resolution
    return math.ceil(width / patch_size) * math.ceil(height / patch_size)


def mb_to_gb(value_mb: float) -> float:
    return value_mb / 1024.0


def pct_reduction(old: float, new: float) -> float:
    if old == 0:
        return 0.0
    return (old - new) / old * 100.0


def get_hardware_check() -> dict:
    check = {
        "checked_at": dt.datetime.now().isoformat(timespec="seconds"),
        "nvidia_smi": None,
        "torch": None,
        "cuda_sanity_allocation": None,
        "errors": []
    }

    try:
        proc = subprocess.run(
            [
                "nvidia-smi",
                "--query-gpu=name,memory.total,driver_version",
                "--format=csv,noheader",
            ],
            capture_output=True,
            text=True,
            timeout=10,
            check=True,
        )
        check["nvidia_smi"] = proc.stdout.strip()
    except Exception as exc:
        check["errors"].append(f"nvidia-smi check failed: {type(exc).__name__}: {exc}")

    try:
        import torch

        torch_info = {
            "torch_version": torch.__version__,
            "cuda_available": bool(torch.cuda.is_available()),
        }
        if torch.cuda.is_available():
            torch_info["device_name"] = torch.cuda.get_device_name(0)
            props = torch.cuda.get_device_properties(0)
            torch_info["total_memory_mb"] = round(props.total_memory / 1024**2, 2)
            free, total = torch.cuda.mem_get_info(0)
            torch_info["free_memory_mb_before"] = round(free / 1024**2, 2)
            torch_info["sanity_note"] = "CUDA device was detected; allocation stress test was skipped to keep the report run non-invasive."
        check["torch"] = torch_info
    except Exception as exc:
        check["errors"].append(f"torch CUDA check failed: {type(exc).__name__}: {exc}")

    write_json(RESULTS_DIR / "hardware_check.json", check)
    return check


def token_memory_mb(tokens: int, config: dict) -> float:
    token_model = config["token_model"]
    return tokens * (
        token_model["activation_mb_per_visual_token"]
        + token_model["kv_cache_mb_per_visual_token"]
    )


def latency_proxy_ms(tokens: int, config: dict, load_count: int, evict_count: int, load_ms: float, evict_ms: float) -> float:
    return (
        config["memory_model"]["base_latency_ms"]
        + tokens * config["token_model"]["latency_ms_per_visual_token"]
        + load_count * load_ms
        + evict_count * evict_ms
    )


def reserve_pass(peak_mb: float, reserve_target_mb: int, config: dict) -> bool:
    return peak_mb + reserve_target_mb <= config["hardware_assumption"]["runtime_budget_mb"]


def compute_peak(base_model_mb: int, tokens: int, adapter_memory_mb: float, temp_buffer_mb: float, config: dict) -> float:
    return base_model_mb + token_memory_mb(tokens, config) + adapter_memory_mb + temp_buffer_mb


def simulate_row(
    baseline: str,
    profile: dict,
    reserve_target_mb: int,
    full_res: list[int],
    global_res: list[int],
    roi_res: list[int],
    roi_count: int,
    adapter_count: int,
    top_k: int,
    config: dict,
    adapter_table: dict,
) -> dict:
    patch_size = config["token_model"]["patch_size"]
    base_model_mb = profile["base_model_memory_mb"]
    full_tokens = tokens_for_resolution(full_res, patch_size)
    global_tokens = tokens_for_resolution(global_res, patch_size)
    roi_tokens_single = tokens_for_resolution(roi_res, patch_size)

    rank = adapter_table.get("simulation_rank", "rank8")
    independent_mem = adapter_table["independent"][f"{rank}_memory_mb"]
    shared_common = adapter_table["shared_hydra_style"]["shared_common_memory_mb"]
    branch_mem = adapter_table["shared_hydra_style"][f"{rank}_branch_memory_mb"]
    temp_base = config["memory_model"]["temporary_load_buffer_mb"]
    min_temp = config["memory_model"]["minimum_temporary_load_buffer_mb"]

    tokens = 0
    resident_adapter_mb = 0.0
    temp_buffer_mb = 0.0
    load_count = 0
    evict_count = 0
    hold_count = 0
    effective_roi_count = roi_count
    effective_top_k = min(top_k, adapter_count)
    adapter_structure = "none"
    policy_trace: list[str] = []

    if baseline == "B0":
        tokens = global_tokens
        effective_roi_count = 0
        effective_top_k = 0
        load_ms = 0
        evict_ms = 0
    elif baseline == "B1":
        tokens = full_tokens
        effective_roi_count = 0
        effective_top_k = 0
        load_ms = 0
        evict_ms = 0
    elif baseline == "B2":
        tokens = global_tokens + roi_tokens_single * roi_count
        effective_top_k = 0
        load_ms = 0
        evict_ms = 0
    elif baseline == "B4-lite":
        tokens = global_tokens + roi_tokens_single * roi_count
        adapter_structure = "independent_all_resident"
        resident_adapter_mb = adapter_count * independent_mem
        temp_buffer_mb = min_temp
        load_count = min(top_k, adapter_count)
        load_ms = adapter_table["independent"]["load_latency_ms"]
        evict_ms = adapter_table["independent"]["evict_latency_ms"]
    elif baseline == "B5-lite":
        tokens = global_tokens + roi_tokens_single * roi_count
        adapter_structure = "shared_all_resident"
        resident_adapter_mb = shared_common + adapter_count * branch_mem
        temp_buffer_mb = min_temp
        load_count = min(top_k, adapter_count)
        load_ms = adapter_table["shared_hydra_style"]["load_latency_ms"]
        evict_ms = adapter_table["shared_hydra_style"]["evict_latency_ms"]
    elif baseline == "B7-lite":
        adapter_structure = "shared_budget_conditioned"
        load_ms = adapter_table["shared_hydra_style"]["load_latency_ms"]
        evict_ms = adapter_table["shared_hydra_style"]["evict_latency_ms"]
        current_roi = roi_count
        current_top_k = min(top_k, adapter_count)
        held_adapter = False

        while True:
            current_tokens = global_tokens + roi_tokens_single * current_roi
            current_adapter_mb = 0.0 if held_adapter or current_top_k == 0 else shared_common + current_top_k * branch_mem
            current_temp = 0.0 if current_adapter_mb == 0 else max(min_temp, min(temp_base, current_top_k * branch_mem))
            current_peak = compute_peak(base_model_mb, current_tokens, current_adapter_mb, current_temp, config)
            if reserve_pass(current_peak, reserve_target_mb, config):
                tokens = current_tokens
                resident_adapter_mb = current_adapter_mb
                temp_buffer_mb = current_temp
                effective_roi_count = current_roi
                effective_top_k = 0 if held_adapter else current_top_k
                break

            if current_top_k > 1:
                current_top_k -= 1
                evict_count += 1
                policy_trace.append("reduce_top_k")
                continue
            if current_roi > 1:
                current_roi -= 1
                policy_trace.append("reduce_roi_count")
                continue
            if not held_adapter:
                held_adapter = True
                current_top_k = 0
                hold_count += 1
                policy_trace.append("hold_adapter_load")
                continue

            tokens = current_tokens
            resident_adapter_mb = current_adapter_mb
            temp_buffer_mb = current_temp
            effective_roi_count = current_roi
            effective_top_k = 0
            policy_trace.append("budget_violation_unresolved")
            break
        load_count = effective_top_k
        if not policy_trace:
            policy_trace.append("allow_initial_plan")
    else:
        raise ValueError(f"Unknown baseline: {baseline}")

    peak_mb = compute_peak(base_model_mb, tokens, resident_adapter_mb, temp_buffer_mb, config)
    passed = reserve_pass(peak_mb, reserve_target_mb, config)
    latency = latency_proxy_ms(tokens, config, load_count, evict_count, load_ms, evict_ms)

    return {
        "run_id": "",
        "baseline": baseline,
        "model_profile": profile["id"],
        "model_name": profile["name"],
        "base_model_memory_mb": base_model_mb,
        "reserve_target_mb": reserve_target_mb,
        "full_res": f"{full_res[0]}x{full_res[1]}",
        "global_res": f"{global_res[0]}x{global_res[1]}",
        "roi_res": f"{roi_res[0]}x{roi_res[1]}",
        "roi_count_requested": roi_count,
        "roi_count_effective": effective_roi_count,
        "adapter_count": adapter_count,
        "active_top_k_requested": top_k,
        "active_top_k_effective": effective_top_k,
        "adapter_structure": adapter_structure,
        "visual_tokens": tokens,
        "visual_memory_mb": round(token_memory_mb(tokens, config), 2),
        "resident_adapter_memory_mb": round(resident_adapter_mb, 2),
        "temporary_load_buffer_mb": round(temp_buffer_mb, 2),
        "estimated_peak_memory_mb": round(peak_mb, 2),
        "estimated_peak_memory_gb": round(mb_to_gb(peak_mb), 3),
        "reserve_pass": passed,
        "reserve_headroom_mb": round(config["hardware_assumption"]["runtime_budget_mb"] - reserve_target_mb - peak_mb, 2),
        "load_count": load_count,
        "evict_count": evict_count,
        "hold_count": hold_count,
        "latency_proxy_ms": round(latency, 2),
        "policy_trace": " > ".join(policy_trace) if policy_trace else "n/a",
    }


def generate_results(config: dict, adapter_table: dict) -> list[dict]:
    rows: list[dict] = []
    baselines = ["B0", "B1", "B2", "B4-lite", "B5-lite", "B7-lite"]
    full_res_values = config["resolutions"]["full_res"]
    global_res_values = config["resolutions"]["global_res"]
    roi_res_values = config["resolutions"]["roi_res"]
    grid = config["grid"]

    for profile in config["model_profiles"]:
        for reserve_target in grid["reserve_targets_mb"]:
            for baseline in baselines:
                if baseline == "B0":
                    for global_res in global_res_values:
                        row = simulate_row(
                            baseline, profile, reserve_target, full_res_values[-1], global_res, roi_res_values[0], 0, 0, 0, config, adapter_table
                        )
                        rows.append(row)
                elif baseline == "B1":
                    for full_res in full_res_values:
                        row = simulate_row(
                            baseline, profile, reserve_target, full_res, global_res_values[0], roi_res_values[0], 0, 0, 0, config, adapter_table
                        )
                        rows.append(row)
                elif baseline == "B2":
                    for global_res in global_res_values:
                        for roi_res in roi_res_values:
                            for roi_count in grid["roi_counts"]:
                                row = simulate_row(
                                    baseline, profile, reserve_target, full_res_values[-1], global_res, roi_res, roi_count, 0, 0, config, adapter_table
                                )
                                rows.append(row)
                else:
                    for global_res in global_res_values:
                        for roi_res in roi_res_values:
                            for roi_count in grid["roi_counts"]:
                                for adapter_count in grid["adapter_counts"]:
                                    for top_k in grid["active_top_k"]:
                                        row = simulate_row(
                                            baseline, profile, reserve_target, full_res_values[-1], global_res, roi_res, roi_count, adapter_count, top_k, config, adapter_table
                                        )
                                        rows.append(row)

    for index, row in enumerate(rows, start=1):
        row["run_id"] = f"costsim-{index:05d}"
    return rows


def write_results(rows: list[dict]) -> None:
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    jsonl_path = RESULTS_DIR / "results.jsonl"
    with jsonl_path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")

    csv_path = RESULTS_DIR / "summary.csv"
    fieldnames = list(rows[0].keys())
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def aggregate_rows(rows: list[dict]) -> list[dict]:
    groups: dict[tuple[str, str, int], list[dict]] = {}
    for row in rows:
        key = (row["model_profile"], row["baseline"], int(row["reserve_target_mb"]))
        groups.setdefault(key, []).append(row)

    aggregates: list[dict] = []
    for (profile, baseline, reserve), items in sorted(groups.items()):
        pass_rate = sum(1 for item in items if item["reserve_pass"]) / len(items)
        aggregates.append(
            {
                "model_profile": profile,
                "baseline": baseline,
                "reserve_target_mb": reserve,
                "runs": len(items),
                "reserve_pass_rate": round(pass_rate, 4),
                "mean_visual_tokens": round(statistics.mean(item["visual_tokens"] for item in items), 2),
                "mean_peak_memory_mb": round(statistics.mean(item["estimated_peak_memory_mb"] for item in items), 2),
                "mean_adapter_memory_mb": round(statistics.mean(item["resident_adapter_memory_mb"] for item in items), 2),
                "mean_latency_proxy_ms": round(statistics.mean(item["latency_proxy_ms"] for item in items), 2),
            }
        )

    path = RESULTS_DIR / "aggregate_summary.csv"
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(aggregates[0].keys()))
        writer.writeheader()
        writer.writerows(aggregates)
    return aggregates


def find_primary_rows(rows: list[dict], config: dict) -> list[dict]:
    primary = config["primary_case"]
    wanted = []
    for baseline in ["B0", "B1", "B2", "B4-lite", "B5-lite", "B7-lite"]:
        candidates = [
            row for row in rows
            if row["baseline"] == baseline
            and row["model_profile"] == primary["model_profile"]
            and row["reserve_target_mb"] == primary["reserve_target_mb"]
        ]
        if baseline == "B0":
            candidates = [row for row in candidates if row["global_res"] == f"{primary['global_res'][0]}x{primary['global_res'][1]}"]
        elif baseline == "B1":
            candidates = [row for row in candidates if row["full_res"] == f"{primary['full_res'][0]}x{primary['full_res'][1]}"]
        elif baseline == "B2":
            candidates = [
                row for row in candidates
                if row["global_res"] == f"{primary['global_res'][0]}x{primary['global_res'][1]}"
                and row["roi_res"] == f"{primary['roi_res'][0]}x{primary['roi_res'][1]}"
                and row["roi_count_requested"] == primary["roi_count"]
            ]
        else:
            candidates = [
                row for row in candidates
                if row["global_res"] == f"{primary['global_res'][0]}x{primary['global_res'][1]}"
                and row["roi_res"] == f"{primary['roi_res'][0]}x{primary['roi_res'][1]}"
                and row["roi_count_requested"] == primary["roi_count"]
                and row["adapter_count"] == primary["adapter_count"]
                and row["active_top_k_requested"] == primary["active_top_k"]
            ]
        if not candidates:
            raise RuntimeError(f"No primary candidate found for {baseline}")
        wanted.append(candidates[0])

    path = RESULTS_DIR / "primary_case_summary.csv"
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(wanted[0].keys()))
        writer.writeheader()
        writer.writerows(wanted)
    return wanted


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = MALGUN_BOLD if bold and MALGUN_BOLD.exists() else MALGUN
    if path.exists():
        return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_bar_chart(path: Path, title: str, labels: list[str], values: list[float], ylabel: str, color: tuple[int, int, int]) -> None:
    width, height = 1200, 720
    margin_left, margin_right = 120, 60
    margin_top, margin_bottom = 110, 120
    chart_w = width - margin_left - margin_right
    chart_h = height - margin_top - margin_bottom

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(34, bold=True)
    label_font = font(22)
    small_font = font(18)

    draw.text((width / 2, 42), title, fill=(20, 24, 28), font=title_font, anchor="mm")
    draw.line((margin_left, margin_top + chart_h, margin_left + chart_w, margin_top + chart_h), fill=(60, 60, 60), width=2)
    draw.line((margin_left, margin_top, margin_left, margin_top + chart_h), fill=(60, 60, 60), width=2)

    max_value = max(values) * 1.15 if values else 1
    for tick in range(5):
        value = max_value * tick / 4
        y = margin_top + chart_h - (value / max_value) * chart_h
        draw.line((margin_left - 8, y, margin_left + chart_w, y), fill=(225, 229, 233), width=1)
        tick_label = f"{value:.1f}" if max_value <= 10 else f"{value:.0f}"
        draw.text((margin_left - 15, y), tick_label, fill=(80, 84, 88), font=small_font, anchor="rm")

    gap = chart_w / len(values)
    bar_w = min(90, gap * 0.58)
    for i, (label, value) in enumerate(zip(labels, values)):
        x = margin_left + gap * i + gap / 2
        bar_h = (value / max_value) * chart_h
        x0 = x - bar_w / 2
        y0 = margin_top + chart_h - bar_h
        x1 = x + bar_w / 2
        y1 = margin_top + chart_h
        draw.rounded_rectangle((x0, y0, x1, y1), radius=8, fill=color)
        draw.text((x, y0 - 14), f"{value:.1f}", fill=(35, 39, 43), font=small_font, anchor="mm")
        draw.text((x, y1 + 30), label, fill=(35, 39, 43), font=label_font, anchor="mm")

    draw.text((margin_left, margin_top - 28), ylabel, fill=(60, 64, 68), font=label_font, anchor="lm")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def draw_pass_rate_chart(path: Path, aggregates: list[dict], primary_profile: str) -> None:
    selected = [
        row for row in aggregates
        if row["model_profile"] == primary_profile and row["reserve_target_mb"] == 4096
    ]
    order = ["B0", "B1", "B2", "B4-lite", "B5-lite", "B7-lite"]
    by_baseline = {row["baseline"]: row for row in selected}
    labels = [baseline for baseline in order if baseline in by_baseline]
    values = [by_baseline[label]["reserve_pass_rate"] * 100 for label in labels]
    draw_bar_chart(path, "Reserve Pass Rate by Baseline (mid profile, reserve 4GB)", labels, values, "pass rate (%)", (64, 126, 150))


def create_figures(primary_rows: list[dict], aggregates: list[dict], config: dict) -> dict[str, Path]:
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    labels = [row["baseline"] for row in primary_rows]
    peak_gb = [row["estimated_peak_memory_gb"] for row in primary_rows]
    adapter_gb = [mb_to_gb(row["resident_adapter_memory_mb"]) for row in primary_rows]

    peak_path = FIGURES_DIR / "primary_peak_memory_bars.png"
    adapter_path = FIGURES_DIR / "primary_adapter_memory_bars.png"
    pass_path = FIGURES_DIR / "reserve_pass_rate.png"

    draw_bar_chart(peak_path, "Estimated Peak Memory - Primary Case", labels, peak_gb, "GB", (85, 105, 180))
    draw_bar_chart(adapter_path, "Resident Adapter Memory - Primary Case", labels, adapter_gb, "GB", (178, 116, 67))
    draw_pass_rate_chart(pass_path, aggregates, config["primary_case"]["model_profile"])

    return {
        "peak": peak_path,
        "adapter": adapter_path,
        "pass_rate": pass_path,
    }


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    out = []
    out.append("| " + " | ".join(headers) + " |")
    out.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        out.append("| " + " | ".join(row) + " |")
    return "\n".join(out)


def build_interpretation(primary_rows: list[dict], aggregates: list[dict], config: dict, all_rows: list[dict]) -> dict:
    by = {row["baseline"]: row for row in primary_rows}
    b1 = by["B1"]
    b2 = by["B2"]
    b4 = by["B4-lite"]
    b5 = by["B5-lite"]
    b7 = by["B7-lite"]

    peak_reduction_b2 = pct_reduction(b1["estimated_peak_memory_mb"], b2["estimated_peak_memory_mb"])
    peak_reduction_b7 = pct_reduction(b1["estimated_peak_memory_mb"], b7["estimated_peak_memory_mb"])
    token_reduction_b2 = pct_reduction(b1["visual_tokens"], b2["visual_tokens"])
    adapter_reduction_b5 = pct_reduction(b4["resident_adapter_memory_mb"], b5["resident_adapter_memory_mb"])
    adapter_reduction_b7 = pct_reduction(b4["resident_adapter_memory_mb"], b7["resident_adapter_memory_mb"])

    primary_profile = config["primary_case"]["model_profile"]
    pass_rates = {
        row["baseline"]: row["reserve_pass_rate"]
        for row in aggregates
        if row["model_profile"] == primary_profile and row["reserve_target_mb"] == 4096
    }
    b7_all = [row for row in all_rows if row["baseline"] == "B7-lite"]
    b4_all = [row for row in all_rows if row["baseline"] == "B4-lite"]
    b7_adjusted = sum(1 for row in b7_all if row["policy_trace"] != "allow_initial_plan")
    b7_hold = sum(1 for row in b7_all if "hold_adapter_load" in row["policy_trace"])
    b4_fail = sum(1 for row in b4_all if not row["reserve_pass"])
    b7_fail = sum(1 for row in b7_all if not row["reserve_pass"])

    return {
        "token_reduction_b2": round(token_reduction_b2, 2),
        "peak_reduction_b2": round(peak_reduction_b2, 2),
        "peak_reduction_b7": round(peak_reduction_b7, 2),
        "adapter_reduction_b5": round(adapter_reduction_b5, 2),
        "adapter_reduction_b7": round(adapter_reduction_b7, 2),
        "pass_rates": pass_rates,
        "primary_trace": b7["policy_trace"],
        "b7_total_runs": len(b7_all),
        "b7_adjusted_runs": b7_adjusted,
        "b7_hold_runs": b7_hold,
        "b4_fail_runs": b4_fail,
        "b7_fail_runs": b7_fail,
    }


def build_report_markdown(primary_rows: list[dict], aggregates: list[dict], hardware: dict, interpretation: dict, figures: dict[str, Path]) -> str:
    summary_rows = []
    for row in primary_rows:
        summary_rows.append(
            [
                row["baseline"],
                str(row["visual_tokens"]),
                f"{row['estimated_peak_memory_gb']:.2f}",
                f"{mb_to_gb(row['resident_adapter_memory_mb']):.2f}",
                "pass" if row["reserve_pass"] else "fail",
                str(row["load_count"]),
                str(row["evict_count"]),
                f"{row['latency_proxy_ms']:.1f}",
            ]
        )

    pass_rate_rows = []
    for row in aggregates:
        if row["model_profile"] == "mid_7b_vlm" and row["reserve_target_mb"] == 4096:
            pass_rate_rows.append([
                row["baseline"],
                str(row["runs"]),
                f"{row['reserve_pass_rate'] * 100:.1f}%",
                f"{mb_to_gb(row['mean_peak_memory_mb']):.2f}",
                f"{mb_to_gb(row['mean_adapter_memory_mb']):.2f}",
            ])

    lines = [
        "# RTX 3090 기반 VRAM Cost Simulation 결과 보고서",
        "",
        f"작성일: {dt.date.today().isoformat()}",
        "",
        "## 1. 요약",
        "",
        "이번 파일럿은 실제 VLM 학습 없이 숫자 기반 cost simulator로 `Budget-Conditioned Foveated Adapter Policy`의 비용 절감 가능성을 확인했다.",
        "",
        f"- B1 full-res 대비 B2 foveated ROI는 visual token을 {interpretation['token_reduction_b2']}% 줄였다.",
        f"- 대표 조건에서 B1 대비 B7-lite의 estimated peak memory 감소율은 {interpretation['peak_reduction_b7']}%였다.",
        f"- B4-lite independent bank 대비 B5-lite shared bank의 resident adapter memory 감소율은 {interpretation['adapter_reduction_b5']}%였다.",
        f"- B4-lite 대비 B7-lite budget policy의 resident adapter memory 감소율은 {interpretation['adapter_reduction_b7']}%였다.",
        f"- 전체 grid에서 B4-lite는 reserve fail {interpretation['b4_fail_runs']}건을 보였고, B7-lite는 fail {interpretation['b7_fail_runs']}건으로 유지됐다.",
        f"- B7-lite는 {interpretation['b7_total_runs']}건 중 {interpretation['b7_adjusted_runs']}건에서 top-k/ROI/adapter hold 조정을 수행했다.",
        "",
        "따라서 이 구조는 최소한 비용 모델 상에서 `full-res 입력 비용 절감`과 `adapter resident memory 절감`이라는 두 축을 분리해서 보여줄 수 있다.",
        "",
        "## 2. 실행 환경 확인",
        "",
        "```json",
        json.dumps(hardware, ensure_ascii=False, indent=2),
        "```",
        "",
        "## 3. 대표 조건 결과표",
        "",
        markdown_table(
            ["baseline", "visual tokens", "peak GB", "adapter GB", "reserve", "load", "evict", "latency proxy ms"],
            summary_rows,
        ),
        "",
        "## 4. Reserve Pass Rate 요약",
        "",
        markdown_table(
            ["baseline", "runs", "pass rate", "mean peak GB", "mean adapter GB"],
            pass_rate_rows,
        ),
        "",
        "## 5. 그래프",
        "",
        f"![Estimated Peak Memory]({figures['peak'].as_posix()})",
        "",
        f"![Resident Adapter Memory]({figures['adapter'].as_posix()})",
        "",
        f"![Reserve Pass Rate]({figures['pass_rate'].as_posix()})",
        "",
        "## 6. 해석",
        "",
        "1차 결과는 foveated ROI 구조가 full-resolution 기준선보다 visual token과 peak memory를 줄일 수 있음을 보여준다. 특히 B2는 adapter를 전혀 쓰지 않아도 입력 비용 절감 축을 분리해서 확인할 수 있다.",
        "",
        "adapter 측면에서는 independent bank를 모두 상주시킨 B4-lite가 adapter 수에 따라 resident memory가 빠르게 커진다. 반면 B5-lite는 shared common과 branch 구조를 가정하기 때문에 같은 adapter count에서도 resident memory 증가율이 낮다.",
        "",
        f"B7-lite의 대표 policy trace는 `{interpretation['primary_trace']}`이다. 대표 조건에서는 초기 계획이 budget 안에 들어왔지만, 전체 grid에서는 {interpretation['b7_adjusted_runs']}건에서 reduce_top_k, reduce_roi_count, hold_adapter_load 같은 조정 행동이 실제로 기록되었다.",
        "",
        "## 7. 한계",
        "",
        "- 이 결과는 실제 VLM profiler 결과가 아니라 cost model 기반 추정이다.",
        "- adapter memory 값은 실제 target model과 LoRA rank가 정해지면 파라미터 수 기반으로 다시 계산해야 한다.",
        "- latency는 실제 커널 실행 시간이 아니라 token 수와 load/evict 횟수 기반 proxy다.",
        "- 정확도는 아직 측정하지 않았으므로, 이 보고서는 성능 우위가 아니라 비용 구조 검증 자료로 해석해야 한다.",
        "",
        "## 8. 다음 단계",
        "",
        "1. 실제 target VLM 후보를 정하고 parameter-count 기반 adapter memory table로 교체한다.",
        "2. 작은 이미지 샘플에서 PyTorch CUDA profiler 또는 실제 inference peak memory를 측정한다.",
        "3. budget policy가 top-k, ROI count, adapter load를 줄이는 stress case를 학교 제출용 그림으로 정리한다.",
        "4. 장비 지원을 받으면 실제 VLM inference, LoRA loading, co-resident module 검증으로 확장한다.",
        "",
        "## 9. 결론",
        "",
        "RTX 3090 24GB 기반 0차 cost simulation은 제안 구조가 VRAM 절약 측면에서 검증할 가치가 있음을 보여준다. 특히 입력 비용 절감과 adapter resident memory 절감을 분리해 측정할 수 있으므로, 후속 장비 지원을 요청하기 위한 초기 근거 자료로 사용할 수 있다.",
        "",
    ]
    return "\n".join(lines)


def set_docx_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_docx_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_docx_font(run, bold=bold)


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], "E8EEF7")
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                set_docx_font(run, bold=True)
        header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    set_docx_font(run, size=9)


def create_docx_report(primary_rows: list[dict], aggregates: list[dict], hardware: dict, interpretation: dict, figures: dict[str, Path]) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RTX 3090 기반 VRAM Cost Simulation 결과 보고서")
    set_docx_font(run, size=18, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"작성일: {dt.date.today().isoformat()}")
    set_docx_font(run, size=10)

    add_docx_paragraph(doc, "1. 요약", "Heading 1")
    bullets = [
        f"B1 full-res 대비 B2 foveated ROI는 visual token을 {interpretation['token_reduction_b2']}% 줄였다.",
        f"대표 조건에서 B1 대비 B7-lite의 estimated peak memory 감소율은 {interpretation['peak_reduction_b7']}%였다.",
        f"B4-lite independent bank 대비 B5-lite shared bank의 resident adapter memory 감소율은 {interpretation['adapter_reduction_b5']}%였다.",
        f"B4-lite 대비 B7-lite budget policy의 resident adapter memory 감소율은 {interpretation['adapter_reduction_b7']}%였다.",
        f"전체 grid에서 B4-lite는 reserve fail {interpretation['b4_fail_runs']}건을 보였고, B7-lite는 fail {interpretation['b7_fail_runs']}건으로 유지됐다.",
        f"B7-lite는 {interpretation['b7_total_runs']}건 중 {interpretation['b7_adjusted_runs']}건에서 top-k/ROI/adapter hold 조정을 수행했다.",
    ]
    for item in bullets:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_docx_font(run)

    add_docx_paragraph(doc, "2. 실행 환경", "Heading 1")
    hw_lines = [
        f"nvidia-smi: {hardware.get('nvidia_smi')}",
        f"torch: {hardware.get('torch')}",
    ]
    for line in hw_lines:
        add_docx_paragraph(doc, line)

    add_docx_paragraph(doc, "3. 대표 조건 결과표", "Heading 1")
    table_rows = [
        [
            row["baseline"],
            str(row["visual_tokens"]),
            f"{row['estimated_peak_memory_gb']:.2f}",
            f"{mb_to_gb(row['resident_adapter_memory_mb']):.2f}",
            "pass" if row["reserve_pass"] else "fail",
            f"{row['latency_proxy_ms']:.1f}",
        ]
        for row in primary_rows
    ]
    add_docx_table(doc, ["baseline", "tokens", "peak GB", "adapter GB", "reserve", "latency ms"], table_rows)

    add_docx_paragraph(doc, "4. 그래프", "Heading 1")
    for caption, image_path in [
        ("Estimated Peak Memory", figures["peak"]),
        ("Resident Adapter Memory", figures["adapter"]),
        ("Reserve Pass Rate", figures["pass_rate"]),
    ]:
        add_docx_paragraph(doc, caption, bold=True)
        doc.add_picture(str(image_path), width=Inches(6.2))

    add_docx_paragraph(doc, "5. 해석과 한계", "Heading 1")
    paragraphs = [
        "1차 결과는 foveated ROI 구조가 full-resolution 기준선보다 visual token과 peak memory를 줄일 수 있음을 보여준다.",
        "adapter 측면에서는 independent bank를 모두 상주시킨 B4-lite가 adapter 수에 따라 resident memory가 빠르게 커지는 반면, shared 구조를 둔 B5-lite와 budget policy를 둔 B7-lite는 더 낮은 resident memory를 보인다.",
        "이 결과는 실제 VLM profiler 결과가 아니라 cost model 기반 추정이다. 따라서 성능 우위가 아니라 비용 구조 검증 자료로 해석해야 한다.",
    ]
    for text in paragraphs:
        add_docx_paragraph(doc, text)

    add_docx_paragraph(doc, "6. 결론", "Heading 1")
    add_docx_paragraph(doc, "RTX 3090 24GB 기반 0차 cost simulation은 제안 구조가 VRAM 절약 측면에서 검증할 가치가 있음을 보여준다. 후속 장비 지원을 요청하기 위한 초기 근거 자료로 사용할 수 있다.")

    output_path = REPORTS_DIR / "results_report_ko.docx"
    doc.save(output_path)
    return output_path


def register_pdf_fonts() -> str:
    if MALGUN.exists():
        pdfmetrics.registerFont(TTFont("Malgun", str(MALGUN)))
        if MALGUN_BOLD.exists():
            pdfmetrics.registerFont(TTFont("MalgunBold", str(MALGUN_BOLD)))
        return "Malgun"
    return "Helvetica"


def create_pdf_report(primary_rows: list[dict], aggregates: list[dict], hardware: dict, interpretation: dict, figures: dict[str, Path]) -> Path:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    font_name = register_pdf_fonts()
    bold_font = "MalgunBold" if font_name == "Malgun" and MALGUN_BOLD.exists() else font_name
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="KTitle", fontName=bold_font, fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=16))
    styles.add(ParagraphStyle(name="KH1", fontName=bold_font, fontSize=13, leading=18, spaceBefore=12, spaceAfter=8))
    styles.add(ParagraphStyle(name="KBody", fontName=font_name, fontSize=9.5, leading=14, alignment=TA_LEFT, spaceAfter=6))
    styles.add(ParagraphStyle(name="KSmall", fontName=font_name, fontSize=8, leading=11, alignment=TA_LEFT))

    output_path = REPORTS_DIR / "results_report_ko.pdf"
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
    )

    story = [
        Paragraph("RTX 3090 기반 VRAM Cost Simulation 결과 보고서", styles["KTitle"]),
        Paragraph(f"작성일: {dt.date.today().isoformat()}", styles["KBody"]),
        Paragraph("1. 요약", styles["KH1"]),
        Paragraph(f"B1 full-res 대비 B2 foveated ROI는 visual token을 {interpretation['token_reduction_b2']}% 줄였다.", styles["KBody"]),
        Paragraph(f"대표 조건에서 B1 대비 B7-lite의 estimated peak memory 감소율은 {interpretation['peak_reduction_b7']}%였다.", styles["KBody"]),
        Paragraph(f"B4-lite independent bank 대비 B5-lite shared bank의 resident adapter memory 감소율은 {interpretation['adapter_reduction_b5']}%였다.", styles["KBody"]),
        Paragraph(f"B4-lite 대비 B7-lite budget policy의 resident adapter memory 감소율은 {interpretation['adapter_reduction_b7']}%였다.", styles["KBody"]),
        Paragraph(f"전체 grid에서 B4-lite는 reserve fail {interpretation['b4_fail_runs']}건을 보였고, B7-lite는 fail {interpretation['b7_fail_runs']}건으로 유지됐다.", styles["KBody"]),
        Paragraph(f"B7-lite는 {interpretation['b7_total_runs']}건 중 {interpretation['b7_adjusted_runs']}건에서 top-k/ROI/adapter hold 조정을 수행했다.", styles["KBody"]),
        Paragraph("2. 실행 환경", styles["KH1"]),
        Paragraph(f"nvidia-smi: {hardware.get('nvidia_smi')}", styles["KBody"]),
        Paragraph(f"torch: {json.dumps(hardware.get('torch'), ensure_ascii=False)}", styles["KSmall"]),
        Paragraph("3. 대표 조건 결과표", styles["KH1"]),
    ]

    table_data = [["baseline", "tokens", "peak GB", "adapter GB", "reserve", "latency ms"]]
    for row in primary_rows:
        table_data.append([
            row["baseline"],
            str(row["visual_tokens"]),
            f"{row['estimated_peak_memory_gb']:.2f}",
            f"{mb_to_gb(row['resident_adapter_memory_mb']):.2f}",
            "pass" if row["reserve_pass"] else "fail",
            f"{row['latency_proxy_ms']:.1f}",
        ])
    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTNAME", (0, 0), (-1, 0), bold_font),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF7")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B9C2CF")),
        ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    story.append(table)
    story.append(PageBreak())
    story.append(Paragraph("4. 그래프", styles["KH1"]))
    for index, (caption, image_path) in enumerate([
        ("Estimated Peak Memory", figures["peak"]),
        ("Resident Adapter Memory", figures["adapter"]),
        ("Reserve Pass Rate", figures["pass_rate"]),
    ]):
        if index == 2:
            story.append(PageBreak())
        story.append(Paragraph(caption, styles["KBody"]))
        story.append(RLImage(str(image_path), width=16.5 * cm, height=9.9 * cm))
        story.append(Spacer(1, 6))

    story.append(PageBreak())
    story.append(Paragraph("5. 해석과 한계", styles["KH1"]))
    story.append(Paragraph("1차 결과는 foveated ROI 구조가 full-resolution 기준선보다 visual token과 peak memory를 줄일 수 있음을 보여준다.", styles["KBody"]))
    story.append(Paragraph("adapter 측면에서는 independent bank를 모두 상주시킨 B4-lite가 adapter 수에 따라 resident memory가 빠르게 커지는 반면, shared 구조를 둔 B5-lite와 budget policy를 둔 B7-lite는 더 낮은 resident memory를 보인다.", styles["KBody"]))
    story.append(Paragraph("이 결과는 실제 VLM profiler 결과가 아니라 cost model 기반 추정이다. 따라서 성능 우위가 아니라 비용 구조 검증 자료로 해석해야 한다.", styles["KBody"]))
    story.append(Paragraph("6. 결론", styles["KH1"]))
    story.append(Paragraph("RTX 3090 24GB 기반 0차 cost simulation은 제안 구조가 VRAM 절약 측면에서 검증할 가치가 있음을 보여준다. 후속 장비 지원을 요청하기 위한 초기 근거 자료로 사용할 수 있다.", styles["KBody"]))

    doc.build(story)
    return output_path


def write_markdown_report(markdown: str) -> Path:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    path = REPORTS_DIR / "results_report_ko.md"
    path.write_text(markdown, encoding="utf-8")
    return path


def render_pdf_preview(pdf_path: Path) -> dict:
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    prefix = RENDER_DIR / "results_report_page"
    result = {
        "pdf": str(pdf_path),
        "rendered": False,
        "pages": [],
        "error": None,
    }
    try:
        subprocess.run(["pdftoppm", "-png", str(pdf_path), str(prefix)], check=True, capture_output=True, text=True, timeout=60)
        pages = sorted(RENDER_DIR.glob("results_report_page-*.png"))
        result["rendered"] = bool(pages)
        result["pages"] = [str(page) for page in pages]
    except Exception as exc:
        result["error"] = f"{type(exc).__name__}: {exc}"
    write_json(RENDER_DIR / "pdf_render_check.json", result)
    return result


def verify_docx_text(docx_path: Path) -> dict:
    result = {
        "docx": str(docx_path),
        "readable": False,
        "paragraph_count": 0,
        "table_count": 0,
        "first_paragraph": "",
        "error": None,
    }
    try:
        doc = Document(str(docx_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        result["readable"] = True
        result["paragraph_count"] = len(paragraphs)
        result["table_count"] = len(doc.tables)
        result["first_paragraph"] = paragraphs[0] if paragraphs else ""
    except Exception as exc:
        result["error"] = f"{type(exc).__name__}: {exc}"
    write_json(RENDER_DIR / "docx_text_check.json", result)
    return result


def main() -> int:
    for directory in [RESULTS_DIR, FIGURES_DIR, REPORTS_DIR, RENDER_DIR]:
        directory.mkdir(parents=True, exist_ok=True)

    config = load_json(CONFIG_DIR / "simulation_config.json")
    adapter_table = load_json(CONFIG_DIR / "adapter_cost_table.json")
    hardware = get_hardware_check()

    rows = generate_results(config, adapter_table)
    write_results(rows)
    aggregates = aggregate_rows(rows)
    primary_rows = find_primary_rows(rows, config)
    figures = create_figures(primary_rows, aggregates, config)
    interpretation = build_interpretation(primary_rows, aggregates, config, rows)

    markdown = build_report_markdown(primary_rows, aggregates, hardware, interpretation, figures)
    markdown_path = write_markdown_report(markdown)
    docx_path = create_docx_report(primary_rows, aggregates, hardware, interpretation, figures)
    pdf_path = create_pdf_report(primary_rows, aggregates, hardware, interpretation, figures)

    pdf_render = render_pdf_preview(pdf_path)
    docx_check = verify_docx_text(docx_path)

    manifest = {
        "created_at": dt.datetime.now().isoformat(timespec="seconds"),
        "root": str(PACKAGE_ROOT),
        "result_rows": len(rows),
        "files": {
            "results_jsonl": str(RESULTS_DIR / "results.jsonl"),
            "summary_csv": str(RESULTS_DIR / "summary.csv"),
            "aggregate_summary_csv": str(RESULTS_DIR / "aggregate_summary.csv"),
            "primary_case_summary_csv": str(RESULTS_DIR / "primary_case_summary.csv"),
            "hardware_check": str(RESULTS_DIR / "hardware_check.json"),
            "peak_figure": str(figures["peak"]),
            "adapter_figure": str(figures["adapter"]),
            "pass_rate_figure": str(figures["pass_rate"]),
            "markdown_report": str(markdown_path),
            "docx_report": str(docx_path),
            "pdf_report": str(pdf_path),
        },
        "checks": {
            "pdf_render": pdf_render,
            "docx_text": docx_check,
        },
        "interpretation": interpretation,
    }
    write_json(RUN_DIR / "manifest.json", manifest)

    print(json.dumps({
        "status": "ok",
        "rows": len(rows),
        "markdown": str(markdown_path),
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        "pdf_pages_rendered": len(pdf_render.get("pages", [])),
        "docx_readable": docx_check.get("readable", False),
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
